from flask import Flask, render_template, request, redirect, url_for, session, flash,send_file,jsonify,make_response
import mysql.connector
import base64
import io
import os
import random
import string
import base64
from docx import Document
from docx.shared import Pt, RGBColor, Inches,Cm
from werkzeug.utils import secure_filename
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from docx2pdf import convert
import pythoncom
from datetime import datetime
from lxml import etree

nsmap = {
    'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
    'a': 'http://schemas.openxmlformats.org/drawingml/2006/main',
    'r': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships',
    'w14': 'http://schemas.microsoft.com/office/word/2010/wordml',
    'wp': 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing',
    'm': 'http://schemas.openxmlformats.org/officeDocument/2006/math',
}

db_connection = mysql.connector.connect(
    host="127.0.0.1",
    user="root",
    password="",
    database="capstoneproject"
)



app = Flask(__name__)
app.secret_key = 'your_secret_key'  # Set a secret key for session management

pdfkit_options = {
    'page-size': 'A4',
    'encoding': 'UTF-8',
}





def replace_placeholder1(doc, placeholder, image_path,font_size=12, alignment=WD_ALIGN_PARAGRAPH.LEFT, bold=False,indentation_spaces=0):
    for paragraph in doc.paragraphs:
        if placeholder in paragraph.text:
            for run in paragraph.runs:
                run.clear()

            # Add the image to the paragraph
            indentation = "               " * indentation_spaces
            paragraph.alignment = alignment
            paragraph.left_indent = Cm(10)
            run = paragraph.add_run()
            run.bold = bold
            run.font.size = Pt(font_size)

            run.add_text(indentation)
            run.add_picture(image_path, width=Cm(2.88), height=Cm(1.56))  # Adjust width and height as needed

def replace_placeholder(doc, placeholder, new_text, font_size=12, bold=False, italic=False, alignment=None):
    for paragraph in doc.paragraphs:
        if placeholder in paragraph.text:
            for run in paragraph.runs:
                if placeholder in run.text:
                    run.text = run.text.replace(placeholder, new_text)
                    run.font.size = Pt(font_size)
                    run.font.bold = bold
                    run.font.italic = italic
                    run.font.color.rgb = RGBColor(0, 0, 0)  # Black color
                    if alignment:
                        run.alignment = alignment

                    

def clear_and_add_line(doc, line_number, new_text, font_size=12, bold=False, italic=False, alignment=None,indentation=None):
    for i, paragraph in enumerate(doc.paragraphs):
        if i == line_number:
            for run in paragraph.runs:
                run.clear()
            run = paragraph.add_run(new_text)
            run.font.size = Pt(font_size)
            run.font.bold = bold
            run.font.italic = italic
            run.font.color.rgb = RGBColor(0, 0, 0)  # Black color
            if alignment:
                run.alignment = alignment

            if indentation is not None:
                # Set the left indentation for the paragraph
                paragraph.paragraph_format.left_indent = Pt(indentation)

            break

def replace_table_cell_placeholder(table, row_index, col_index, new_text):
    cell = table.cell(row_index, col_index)
    
    # Clear the existing text in the cell
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.text = ""
    
    # Add the new text to the cell
    cell.text = new_text

def replace_table_cell_placeholder1(table, row_index, col_index, new_text, placeholder):
    cell = table.cell(row_index, col_index)
    
    
    # Flag to track if the placeholder was found and replaced
    placeholder_replaced = False
    if new_text == "":
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                if placeholder in run.text:
                    
                    run.text = run.text.replace(placeholder, "N/A")
                    placeholder_replaced = True

        # Add the new text to the cell if the placeholder was not found
        if not placeholder_replaced:
            cell.text = new_text
    else:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                if placeholder in run.text:
                    
                    run.text = run.text.replace(placeholder, new_text)
                    placeholder_replaced = True

        # Add the new text to the cell if the placeholder was not found
        if not placeholder_replaced:
            cell.text = new_text

def replace_table_cell_placeholder2(table, row_index, col_index, new_text, placeholder):
    cell = table.cell(row_index, col_index)
    
    
    
    if new_text == "checked":
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                if placeholder in run.text:
                    
                    run.text = run.text.replace(placeholder, "☑")

    else:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                if placeholder in run.text:
                    
                    run.text = run.text.replace(placeholder, "☐")
                    pdfkit_options

                 


def toggle_table_cell_checkbox(table, row_index, col_index, status):
    cell = table.cell(row_index, col_index)
    print(status)
    if status == "checked":
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.text = ""
        cell.text = "☑"

        # Center-align the text horizontally
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(11)  # Adjust font size if needed
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    else:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.text = ""
        cell.text = "☐"

        # Center-align the text horizontally
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(11)  # Adjust font size if needed
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER




def replace_table_cell_placeholder_with_image(table, row_index, col_index, image_path, placeholder, indentation_spaces=0):
    cell = table.cell(row_index, col_index)

    # Iterate through the runs in the cell
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            if placeholder in run.text:
                # Clear the existing run
                run.clear()

                # Add an image with the specified path
                indentation = "   " * indentation_spaces
                paragraph.left_indent = Cm(3)
                run = paragraph.add_run()
                run.add_text(indentation)
                run.add_picture(image_path, width=Cm(2.88), height=Cm(1.56)) 

def generate_random_code(length=8):
    # Define the characters to choose from
    characters = string.ascii_uppercase + string.digits

    # Generate a random code
    code = '#' + ''.join(random.choice(characters) for _ in range(length - 1))

    return code

@app.route('/submit_report', methods=['POST'])
def submit_report():
    pythoncom.CoInitialize()
    
    kind = request.form.get('kind')
    if kind == "Formal Complaint":
        department = request.form.get('department')
        provision = request.form.get('provision')
        final = request.form.get('final')
        report_text = request.form.get('narrate')
        name = request.form.get('name')
        section = request.form.get('section')
        number = request.form.get('number')
        email = request.form.get('email')
        witness1 = request.form.get('witness1')
        witness2 = request.form.get('witness2')
        witness3 = request.form.get('witness3')
        evidence1 = request.form.get('witness1')
        evidence2 = request.form.get('witness2')
        evidence3 = request.form.get('witness3')
        pic = request.files['file2']
        current_datetime = datetime.now()
        current_date = current_datetime.date()
        formatted_date = current_date.strftime("/%m/%d/%Y") 
        random_code = generate_random_code()

        print(pic)

        if department == "CAFAD":
            Name_Coordinator = "CAFAD Coordinator"

        elif department == "CICS":
            Name_Coordinator = "Lovely Rose Tipan Hernandez"
            
            
        username = session.get('username', '')
        print(username)
        
        pdf_filename = 'Formal Complaint Letter.docx'

        doc = Document(pdf_filename)
        # Replace placeholders

        replace_table_cell_placeholder1(doc.tables[0], 2, 2, formatted_date,"(date)")
        replace_table_cell_placeholder1(doc.tables[0], 4, 1, Name_Coordinator,"NAME")
        replace_table_cell_placeholder1(doc.tables[0], 11, 8, name,"(student)")
        replace_table_cell_placeholder1(doc.tables[0], 12, 8, department,"(college)")
        replace_table_cell_placeholder1(doc.tables[0], 13, 8, section,"(section)")
        replace_table_cell_placeholder1(doc.tables[0], 16, 3, provision,"(provision)")
        replace_table_cell_placeholder1(doc.tables[0], 23, 3, report_text,"(narration)")
        replace_table_cell_placeholder1(doc.tables[0], 30, 3, final,"(final)")
        replace_table_cell_placeholder_with_image(doc.tables[0], 37, 18, pic,"lol")
        replace_table_cell_placeholder1(doc.tables[0], 38, 18, number,"(number)")
        replace_table_cell_placeholder1(doc.tables[0], 39, 18, email, "(email)")
        replace_table_cell_placeholder1(doc.tables[0], 40, 6, witness1,"(witness1)")
        replace_table_cell_placeholder1(doc.tables[0], 41, 6, witness2,"(witness2)")
        replace_table_cell_placeholder1(doc.tables[0], 42, 6, witness3,"(witness3)")
        replace_table_cell_placeholder1(doc.tables[0], 44, 9, evidence1,"(evidence1)")
        replace_table_cell_placeholder1(doc.tables[0], 45, 9, evidence2,"(evidence2)")
        replace_table_cell_placeholder1(doc.tables[0], 46, 9, evidence3,"(evidence3)")

        doc.save("modified_document.docx")
         # Convert the Word document to PDF using docx2pdf
        pdf_path = os.path.join('modified_document.pdf')
        convert("modified_document.docx", pdf_path)

    

        file_name = f'{random_code}_Formal Complaint'
        with open(pdf_path, "rb") as pdf_file:
            pdf_data = pdf_file.read()


        
            
        
        # Check if the POST request has the file part for the supporting document file
        if 'file1' not in request.files:
            flash('No supporting document file part')
            return redirect(request.url)
        
        support_file = request.files['file1']
        
        # Check if the user submitted an empty supporting document file input
    
        if support_file.filename == '':
            support_data = None
            support_filename = "None"
            support_extension = "None"

            db_cursor = db_connection.cursor()
            db_cursor.execute("INSERT INTO reports (report_id, course, report, file_form, file_form_name,file_support_name, file_support_type, file_support, username, date_time, status) VALUES (%s,%s, %s, %s, %s, %s, %s, %s, %s, %s, %s)",
                        (random_code, department, report_text,   pdf_data, file_name,support_filename, support_extension, support_data, username, current_datetime, "Pending"))
            db_connection.commit()
            
            db_cursor.close()
            os.remove("modified_document.docx")
            os.remove("modified_document.pdf")
            
            
            flash('The report is submitted', 'success')
            return redirect('/hello')
            
        
        else:
            # Securely get the filenames and file extensions
            support_filename = secure_filename(support_file.filename)
        
            support_extension = os.path.splitext(support_filename)[1]
            
            # Read the file data into memory
            
            support_data = support_file.read()
            
            
            # Insert the report with file information into the database, including file data
            db_cursor = db_connection.cursor()
            db_cursor.execute("INSERT INTO reports (report_id, course, report, file_form, file_form_name,file_support_name, file_support_type, file_support, username, date_time, status) VALUES (%s,%s, %s, %s, %s, %s, %s, %s, %s, %s, %s)",
                        (random_code, department, report_text,   pdf_data, file_name,support_filename, support_extension, support_data, username, current_datetime, "Pending"))
            db_connection.commit()
            
            db_cursor.close()
            os.remove("modified_document.docx")
            os.remove("modified_document.pdf")
            
            
            flash('The report is submitted', 'success')
            return redirect('/hello')
    else:
        department = request.form.get('department')
        remarks = request.form.get('remarks')
        report_text = request.form.get('narrate1')
        name = request.form.get('name1')
        section = request.form.get('section1')
        designation = request.form.get('designation')
        program = request.form.get('program') 
        pic = request.files['file3']
        current_datetime = datetime.now()
        current_date = current_datetime.date()
        formatted_date = current_date.strftime("/%m/%d/%Y") 
        current_time = current_datetime.strftime('%I:%M %p')
        random_code = generate_random_code()
                 
        username = session.get('username', '')

        pdf_filename = 'Incident Report.docx'
        doc = Document(pdf_filename)
        # Replace placeholders
    
        replace_table_cell_placeholder(doc.tables[0], 2, 3, str(current_date))
        replace_table_cell_placeholder(doc.tables[0], 3, 3, name)
        replace_table_cell_placeholder(doc.tables[0], 4, 3, department)
        replace_table_cell_placeholder(doc.tables[0], 5, 3, program)
        replace_table_cell_placeholder(doc.tables[0], 6, 4, report_text)
        replace_table_cell_placeholder(doc.tables[0], 10, 4, remarks)
        
        replace_table_cell_placeholder_with_image(doc.tables[0], 14, 1, pic,"(signature)",29)
        replace_table_cell_placeholder1(doc.tables[0], 14, 1, designation,"(designation)")
        replace_table_cell_placeholder1(doc.tables[0], 14, 1, str(current_date),"lol")
        replace_table_cell_placeholder(doc.tables[0], 2, 8, current_time)
        replace_table_cell_placeholder(doc.tables[0], 3, 10, username)
        replace_table_cell_placeholder(doc.tables[0], 5, 8, section)
        
        replace_table_cell_placeholder(doc.tables[1], 2, 3, str(current_date))
        replace_table_cell_placeholder(doc.tables[1], 3, 3, name)
        replace_table_cell_placeholder(doc.tables[1], 4, 3, department)
        replace_table_cell_placeholder(doc.tables[1], 5, 3, program)
        replace_table_cell_placeholder(doc.tables[1], 6, 4, report_text)
        replace_table_cell_placeholder(doc.tables[1], 10, 4, remarks)
        
        replace_table_cell_placeholder_with_image(doc.tables[1], 14, 1, pic,"(signature)",29)
        replace_table_cell_placeholder1(doc.tables[1], 14, 1, designation,"(designation)")
        replace_table_cell_placeholder1(doc.tables[1], 14, 1, str(current_date),"lol")
        replace_table_cell_placeholder(doc.tables[1], 2, 8, current_time)
        replace_table_cell_placeholder(doc.tables[1], 3, 10, username)
        replace_table_cell_placeholder(doc.tables[1], 5, 8, section)
        
    
        
        doc.save("modified_document.docx")
        pdf_path = os.path.join('modified_document.pdf')
        convert("modified_document.docx", pdf_path)

    

        file_name = f'{random_code}_Incident Report'
        with open(pdf_path, "rb") as pdf_file:
            pdf_data = pdf_file.read()

        
            
        
        # Check if the POST request has the file part for the supporting document file
        if 'file4' not in request.files:
            flash('No supporting document file part')
            return redirect(request.url)
        
        support_file = request.files['file4']
        
        # Check if the user submitted an empty supporting document file input
    
        if support_file.filename == '':
            support_data = None
            support_filename = "None"
            support_extension = "None"

            db_cursor = db_connection.cursor()
            db_cursor.execute("INSERT INTO reports (report_id, course, report, file_form, file_form_name,file_support_name, file_support_type, file_support, username, date_time, status) VALUES (%s,%s, %s, %s, %s, %s, %s, %s, %s, %s, %s)",
                        (random_code, department, report_text,   pdf_data, file_name,support_filename, support_extension, support_data, username, current_datetime, "Pending"))
            db_connection.commit()
            
            db_cursor.close()
            os.remove("modified_document.docx")
            os.remove("modified_document.pdf")
            
            
            flash('The report is submitted', 'success')
            return redirect('/hello')
            
        
        else:
            # Securely get the filenames and file extensions
            support_filename = secure_filename(support_file.filename)
        
            support_extension = os.path.splitext(support_filename)[1]
            
            # Read the file data into memory
            
            support_data = support_file.read()
            
            
            # Insert the report with file information into the database, including file data
            db_cursor = db_connection.cursor()
            db_cursor.execute("INSERT INTO reports (report_id, course, report, file_form, file_form_name,file_support_name, file_support_type, file_support, username, date_time, status) VALUES (%s,%s, %s, %s, %s, %s, %s, %s, %s, %s, %s)",
                        (random_code, department, report_text,   pdf_data, file_name,support_filename, support_extension, support_data, username, current_datetime, "Pending"))
            db_connection.commit()
            
            db_cursor.close()
            os.remove("modified_document.docx")
            os.remove("modified_document.pdf")
            
            
            flash('The report is submitted', 'success')
            return redirect('/hello')

@app.route('/submit_request', methods=['GET', 'POST'])
def submit_request():
    pythoncom.CoInitialize()
    kind = request.form.get('forms')
    print(kind)
    if kind == "Temporary Gate Pass":

        remarks = request.form.get('remarks')
        department = request.form.get('department')
        print(department)
        section = request.form.get('section2')
        program = request.form.get('program')
        current_datetime = datetime.now()
        random_code = generate_random_code()
        current_date = current_datetime.date()
        formatted_date = current_date.strftime("/%m/%d/%Y") 

        student = session.get('namestudent', '') 
        username = session.get('username', '')

        if department == "CAFAD":
            Name_Coordinator1 = "CAFAD Coordinator"

        elif department == "CICS":
            Name_Coordinator1 = "Lovely Rose Tipan Hernandez"

        pdf_filename = 'Temporary Gate Pass.docx'
        doc = Document(pdf_filename)

        

        replace_table_cell_placeholder1(doc.tables[0], 2, 11, formatted_date,"(date)")
        replace_table_cell_placeholder1(doc.tables[0], 11, 1, formatted_date,"(date1)")
        replace_table_cell_placeholder1(doc.tables[0], 11, 1, Name_Coordinator1,"(Coord)")
        replace_table_cell_placeholder1(doc.tables[0], 3, 3, student,"(name)")
        replace_table_cell_placeholder1(doc.tables[0], 6, 4, remarks,"(remarks)")
        replace_table_cell_placeholder1(doc.tables[0], 3, 11, username,"(code)")
        replace_table_cell_placeholder1(doc.tables[0], 5, 9, section, "(section)")
        replace_table_cell_placeholder1(doc.tables[0], 4, 3, department, "(department)")
        replace_table_cell_placeholder1(doc.tables[0], 5, 3, program, "(program)")

        replace_table_cell_placeholder1(doc.tables[1], 2, 11, formatted_date,"(date)")
        replace_table_cell_placeholder1(doc.tables[1], 11, 1, formatted_date,"(date1)")
        replace_table_cell_placeholder1(doc.tables[1], 11, 1, Name_Coordinator1,"(Coord)")
        replace_table_cell_placeholder1(doc.tables[1], 3, 3, student,"(name)")
        replace_table_cell_placeholder1(doc.tables[1], 6, 4, remarks,"(remarks)")
        replace_table_cell_placeholder1(doc.tables[1], 3, 11, username,"(code)")
        replace_table_cell_placeholder1(doc.tables[1], 5, 9, section, "(section)")
        replace_table_cell_placeholder1(doc.tables[1], 4, 3, department, "(department)")
        replace_table_cell_placeholder1(doc.tables[1], 5, 3, program, "(program)")

        doc.save("modified_document.docx")
        pdf_path = os.path.join('modified_document.pdf')
        convert("modified_document.docx", pdf_path)

    

        file_name = f'{random_code}_Temporary Gate Pass'
        with open(pdf_path, "rb") as pdf_file:
            pdf_data = pdf_file.read()
        
        
        if 'file5' not in request.files:
            flash('No supporting document file part')
            return redirect(request.url)
        
        support_file = request.files['file5']
        
        # Check if the user submitted an empty supporting document file input
    
        if support_file.filename == '':
            support_file = None 
        
       
        
        if support_file:
            # Securely get the filenames and file extensions
           
            support_filename = secure_filename(support_file.filename)
            support_extension = os.path.splitext(support_filename)[1]
            
        
            support_data = support_file.read()
            
            # Insert the report with file information into the database, including file data
            db_cursor = db_connection.cursor()
            db_cursor.execute("INSERT INTO forms_osd (form_id,course, report, file_form_name, file_form, file_support_name, file_support_type, file_support, username, date_time, status) VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s)",
                            (random_code, department, remarks, file_name, pdf_data, support_filename, support_extension, support_data, username, current_datetime,"Pending"))
            db_connection.commit()
            db_cursor.close()
            os.remove("modified_document.docx")
            os.remove("modified_document.pdf")
            
            flash('Report submitted successfully')
            return redirect('/hello')
        
    elif kind == "Request for Non-Wearing of Uniform":
        fieldwork = request.form.get('fieldwork')
        prolonged = request.form.get('prolonged')
        foreign = request.form.get('foreign')
        pregnant = request.form.get('pregnant')
        cases = request.form.get('cases')
        majeure = request.form.get('majeure')
        internship = request.form.get('internship')
        specify = request.form.get('specify')
        remarks = request.form.get('remarks')
        department = request.form.get('department')
        specify1 = request.form.get('specifyTextarea')
        print(specify1)
        print(department)
        section = request.form.get('section2')
        college = request.form.get('college')
        program = request.form.get('program')
        current_datetime = datetime.now()
        random_code = generate_random_code()
        current_date = current_datetime.date()
        formatted_date = current_date.strftime("/%m/%d/%Y") 
        pic = request.files['file3']

        student = session.get('namestudent', '') 
        username = session.get('username', '')

        if fieldwork == "fieldwork":
            status = "checked"
        else:
            status = "not"

        if prolonged == "prolonged":
            status1 = "checked"
        else:
            status1 = "not"

        if foreign == "foreign":
            status2 = "checked"
        else:
            status2 = "not"

        if pregnant == "pregnant":
            status3 = "checked"
        else:
            status3 = "not"

        if cases == "cases":
            status4 = "checked"
        else:
            status4 = "not"

        if majeure == "majeure":
            status5 = "checked"
        else:
            status5 = "not"

        if internship == "internship":
            status6 = "checked"
        else:
            status6 = "not"

        if specify == "specify":
            status7 = "checked"
        else:
            status7 = "not"

        if department == "CAFAD":
            Name_Coordinator1 = "CAFAD Coordinator"

        elif department == "CICS":
            Name_Coordinator1 = "Lovely Rose Tipan Hernandez"

        pdf_filename = 'Request for Non-Wearing of Uniform.docx'
        doc = Document(pdf_filename)

        toggle_table_cell_checkbox(doc.tables[0], 6, 0, status)
        toggle_table_cell_checkbox(doc.tables[0], 7, 0, status1)
        toggle_table_cell_checkbox(doc.tables[0], 8, 0, status2)
        toggle_table_cell_checkbox(doc.tables[0], 9, 0, status3)
        toggle_table_cell_checkbox(doc.tables[0], 10, 0, status4)
        toggle_table_cell_checkbox(doc.tables[0], 11, 0, status5)
        toggle_table_cell_checkbox(doc.tables[0], 12, 0, status6)
        toggle_table_cell_checkbox(doc.tables[0], 13, 0, status7)


        replace_table_cell_placeholder1(doc.tables[0], 16, 1, formatted_date,"(date)")
        replace_table_cell_placeholder1(doc.tables[0], 13, 2, specify1,"(specify)")
        replace_table_cell_placeholder1(doc.tables[0], 2, 4, student,"(name)")
        replace_table_cell_placeholder1(doc.tables[0], 3, 4, college,"(college)")
        replace_table_cell_placeholder1(doc.tables[0], 4, 4, program,"(program)")
        replace_table_cell_placeholder1(doc.tables[0], 2, 10, username,"(srcode)")
        replace_table_cell_placeholder1(doc.tables[0], 4, 10, section,"(section)")
        replace_table_cell_placeholder_with_image(doc.tables[0], 16, 1, pic,"(signature)",29)

        




        

        doc.save("modified_document.docx")
        pdf_path = os.path.join('modified_document.pdf')
        convert("modified_document.docx", pdf_path)

    

        file_name = f'{random_code}_Request for Non-Wearing of Uniform'
        with open(pdf_path, "rb") as pdf_file:
            pdf_data = pdf_file.read()
        
        
        if 'file6' not in request.files:
            flash('No supporting document file part')
            return redirect(request.url)
        
        support_file = request.files['file6']
        
        # Check if the user submitted an empty supporting document file input
    
        if support_file.filename == '':
            support_file = None 
        
       
        
        if support_file:
            # Securely get the filenames and file extensions
           
            support_filename = secure_filename(support_file.filename)
            support_extension = os.path.splitext(support_filename)[1]
            
        
            support_data = support_file.read()
            
            # Insert the report with file information into the database, including file data
            db_cursor = db_connection.cursor()
            db_cursor.execute("INSERT INTO forms_osd (form_id,course,file_form_name, file_form, file_support_name, file_support_type, file_support, username, date_time, status) VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s)",
                            (random_code, department,file_name, pdf_data, support_filename, support_extension, support_data, username, current_datetime,"Pending"))
            db_connection.commit()
            db_cursor.close()
            os.remove("modified_document.docx")
            os.remove("modified_document.pdf")
            
            flash('Report submitted successfully')
            return redirect('/hello')

            #Request for new id 
    elif kind == "Request for New ID":
        fieldwork = request.form.get('fieldwork')
        prolonged = request.form.get('prolonged')
        foreign = request.form.get('foreign')
        pregnant = request.form.get('pregnant')    
        specify2 = request.form.get('specify1')
        print(specify2)
        remarks = request.form.get('remarks')
        department = request.form.get('department')
        specify3 = request.form.get('specifyTextarea1')
        section = request.form.get('section1')
        college = request.form.get('college')
        program = request.form.get('program')
        current_datetime = datetime.now()
        random_code = generate_random_code()
        current_date = current_datetime.date()
        formatted_date = current_date.strftime("/%m/%d/%Y") 
        pic = request.files['file8']
        print(pic)

        student = session.get('namestudent', '') 
        username = session.get('username', '')

        if fieldwork == "fieldwork":
            status = "checked"
        else:
            status = "not"

        if prolonged == "prolonged":
            status1 = "checked"
        else:
            status1 = "not"

        if foreign == "foreign":
            status2 = "checked"
        else:
            status2 = "not"

        if pregnant == "pregnant":
            status3 = "checked"

        else:
            status3 = "not"


        if specify2 == "specify1":
            status7 = "checked"
        else:
            status7 = "not"

        if department == "CAFAD":
            Name_Coordinator1 = "CAFAD Coordinator"

        elif department == "CICS":
            Name_Coordinator1 = "Lovely Rose Tipan Hernandez"

        pdf_filename = 'request for new id.docx'
        doc = Document(pdf_filename)
#problem
        replace_table_cell_placeholder2(doc.tables[0], 7, 0, status,"SHIFT")
        replace_table_cell_placeholder2(doc.tables[0], 7, 4, status1,"LOST")
        replace_table_cell_placeholder2(doc.tables[0], 7, 8, status2,"TORN")
        replace_table_cell_placeholder2(doc.tables[0], 9, 2, status3,"UPDATE")
        replace_table_cell_placeholder2(doc.tables[0], 9, 4, status7,"OTHERS")

        
        replace_table_cell_placeholder1(doc.tables[0], 2, 3, formatted_date,"(date)")
        replace_table_cell_placeholder1(doc.tables[0], 10, 1, formatted_date,"(date1)")
        replace_table_cell_placeholder1(doc.tables[0], 10, 8, Name_Coordinator1,"NAME")
        replace_table_cell_placeholder1(doc.tables[0], 8, 8, specify3,"(specify)")
        replace_table_cell_placeholder1(doc.tables[0], 3, 3, student,"(name)")
        replace_table_cell_placeholder1(doc.tables[0], 4, 3, college,"(college)")
        replace_table_cell_placeholder1(doc.tables[0], 5, 3, program,"(program)")
        replace_table_cell_placeholder1(doc.tables[0], 3, 10, username,"(srcode)")
        replace_table_cell_placeholder1(doc.tables[0], 5, 10, section,"(yearlevel)")
        replace_table_cell_placeholder_with_image(doc.tables[0], 10, 1, pic,"(signature)",29)
        
       


        doc.save("modified_document.docx")
        pdf_path = os.path.join('modified_document.pdf')
        convert("modified_document.docx", pdf_path)

    

        file_name = f'{random_code}_Request for New ID'
        with open(pdf_path, "rb") as pdf_file:
            pdf_data = pdf_file.read()
        
        
        if 'file4' not in request.files:
            flash('No supporting document file part')
            return redirect(request.url)
        
        support_file = request.files['file4']
        
        # Check if the user submitted an empty supporting document file input
    
        if support_file.filename == '':
            support_file = None 
        
       
        
        if support_file:
            # Securely get the filenames and file extensions
           
            support_filename = secure_filename(support_file.filename)
            support_extension = os.path.splitext(support_filename)[1]
            
        
            support_data = support_file.read()
            
            # Insert the report with file information into the database, including file data
            db_cursor = db_connection.cursor()
            db_cursor.execute("INSERT INTO forms_osd (form_id,course,file_form_name, file_form, file_support_name, file_support_type, file_support, username, date_time, status) VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s)",
                            (random_code, department,file_name, pdf_data, support_filename, support_extension, support_data, username, current_datetime,"Pending"))
            db_connection.commit()
            db_cursor.close()
            os.remove("modified_document.docx")
            os.remove("modified_document.pdf")
            
            flash('Report submitted successfully')
            return redirect('/hello')


@app.route('/submit_approve', methods=['GET', 'POST'])
def submit_approve():
    remarks = request.form.get('remarks')
    report_id = session.get('Idreport', '') 
    modified_file = request.files['file4']
    support_data = modified_file.read()
    print(report_id)
    print(remarks)
    status = "Approved"

    db_cursor = db_connection.cursor()
    db_cursor.execute("UPDATE forms_osd SET remarks = %s, status = %s, file_form = %s WHERE form_id = %s",(remarks,status,support_data,report_id))
    db_connection.commit()
    db_cursor.close()

    
    return redirect('/request')

@app.route('/submit_reject', methods=['GET', 'POST'])
def submit_reject():
    remarks = request.form.get('remarks')
    report_id = session.get('Idreport', '') 
    print(report_id)
    print(remarks)
    status = "Rejected"

    db_cursor = db_connection.cursor()
    db_cursor.execute("UPDATE forms_osd SET remarks = %s, status = %s WHERE form_id = %s",(remarks,status,report_id))
    db_connection.commit()
    db_cursor.close()

    
    return redirect('/request')

@app.route('/submit_sanction', methods=['POST'])
def submit_sanction():
    if request.method == 'POST':
        # Get the values from the form
        name = session.get('name')
        course = session.get('course')
        sanction = request.form['sanctions']
        # Get the current date and time
        current_datetime = datetime.now()
        # Get the username from the session
       

        # Insert the values into the database
        db_cursor = db_connection.cursor()
        db_cursor.execute("INSERT INTO sanctions (username, course, date_time, sanction) VALUES (%s, %s, %s, %s)",
                          (name, course, current_datetime, sanction))
        db_connection.commit()
        db_cursor.close()

        # Optionally, you can redirect to a success page or perform other actions
        
        return redirect(url_for('homepage'))

    flash('Report submitted successfully!', 'success')
    return redirect(url_for('homepage'))

@app.route('/', methods=['GET', 'POST'])
def index():
    # Retrieve the username from the session if it exists
    username = session.get('username', '')

    error_message = None

    if request.method == 'POST':
        # Get the submitted username and password
        submitted_username = request.form['username']
        submitted_password = request.form['lname']  # Assuming password input is named 'lname'

        # Query the accounts_cics table to check if the provided username and password exist
        db_cursor = db_connection.cursor()
        db_cursor.execute("SELECT * FROM accounts_cics WHERE username = %s AND password = %s",
                          (submitted_username, submitted_password))
        result_cics = db_cursor.fetchone()
        db_cursor.close()

        # Query the account_coordinators table to check if the provided username and password exist
        db_cursor = db_connection.cursor()
        db_cursor.execute("SELECT * FROM accounts_coordinators WHERE username = %s AND password = %s",
                          (submitted_username, submitted_password))
        result_coordinators = db_cursor.fetchone()
        db_cursor.close()

        if result_cics:
            # User exists in the accounts_cics table, set the role and continue
            session['username'] = submitted_username  # Save the username in the session
            session['role'] = 'accounts_cics'  # Save the role in the session
            return redirect(url_for('homepage'))  # Redirect to another page after successful login
        elif result_coordinators:
            # User exists in the accounts_coordinators table, set the role and continue
            session['username'] = submitted_username  # Save the username in the session
            session['role'] = 'accounts_coordinators'  # Save the role in the session
            return redirect(url_for('homepage'))  # Redirect to another page after successful login
        else:
            # Invalid username or password, set error_message
            error_message = 'Invalid username or password'

    return render_template('index.html', username=username, error_message=error_message)


@app.route('/menu')
def menu():
    # Retrieve the username and role from the session
    username = session.get('username', '')
    user_role = session.get('role', '')
    user_source = session.get('source', '')


    # Query the database to retrieve reports for the logged-in user
    db_cursor = db_connection.cursor()

    if user_role == 'accounts_coordinators':
        # If the user is an accounts coordinator, retrieve the course of the user
        db_cursor.execute("SELECT course FROM accounts_coordinators WHERE username = %s", (username,))
        user_course = db_cursor.fetchone()

        if user_course:
            user_course = user_course[0]  # Extract the course from the result

            # Query reports where the course matches the user's course
            db_cursor.execute("SELECT * FROM reports WHERE course = %s", (user_course,))
            reports = db_cursor.fetchall()
    else:
        # For other roles, simply retrieve reports for the logged-in user
        db_cursor.execute("SELECT * FROM reports WHERE username = %s", (username,))
        reports = db_cursor.fetchall()
        user_course = ""

    # Close the cursor
    db_cursor.close()

    return render_template('menu.html', reports=reports, user_source=user_source, user_course=user_course)

@app.route('/request', methods=['GET', 'POST'])
def requestpage():
    # Retrieve the username and role from the session
    username = session.get('username', '')
    user_role = session.get('role', '')
    user_source = session.get('source', '')


    # Query the database to retrieve reports for the logged-in user
    db_cursor = db_connection.cursor()

    if user_role == 'accounts_coordinators':
        # If the user is an accounts coordinator, retrieve the course of the user
        db_cursor.execute("SELECT course FROM accounts_coordinators WHERE username = %s", (username,))
        user_course = db_cursor.fetchone()

        if user_course:
            user_course = user_course[0]  # Extract the course from the result

            # Query reports where the course matches the user's course
            db_cursor.execute("SELECT * FROM forms_osd WHERE course = %s", (user_course,))
            reports = db_cursor.fetchall()
    else:
        # For other roles, simply retrieve reports for the logged-in user
        db_cursor.execute("SELECT * FROM forms_osd WHERE username = %s", (username,))
        reports = db_cursor.fetchall()
        user_course = ""

    # Close the cursor
    db_cursor.close()

    return render_template('request.html', reports=reports, user_source=user_source, user_course=user_course)


@app.route('/save_Id', methods=['GET', 'POST'])
def save_Id():
    report_id = request.form.get('reportId')
    print("Received report_id:", report_id)
    session['Idreport'] = report_id
    return '', 204
        

@app.route('/search_students', methods=['POST'])
def search_students():
    if request.method == 'POST':
        search_value = request.form['username']  # Updated to match the input name
        session['search_value'] = search_value

        

        # Perform a database query to search for students in the accounts_cics table
        db_cursor = db_connection.cursor()
        db_cursor.execute("SELECT * FROM accounts_cics WHERE Name LIKE %s", ('%' + search_value + '%',))
        search_results = db_cursor.fetchall()
        db_cursor.close()

        db_cursor = db_connection.cursor()
        db_cursor.execute("SELECT * FROM sanctions WHERE username LIKE %s", ('%' + search_value + '%',))
        search_results = db_cursor.fetchall()
        db_cursor.close()

        # Check if any results were found
        if search_results:
            # Retrieve the first result (you can modify this logic as needed)
            first_result = search_results[0]

            # Extract the name and course from the result
            name = first_result['Name']
            course = first_result['CourseOrPosition']
          


            # Return the name and course as JSON
            return jsonify({'name': name, 'course': course, 'search_value': session.get('search_value')})
        else:
            # No results found, return an error message as JSON
            print(session.get('name'))
            print(session.get('course'))
            return jsonify({'error': 'No results found'})
        
        
@app.route('/forms')
def forms():
    username = session.get('username', '')
    user_role = session.get('role', '')
    user_source = session.get('source', '')
    print(f"Student Name: {user_role}")

    # Query the database to retrieve reports for the logged-in user
    db_cursor = db_connection.cursor()

    if user_role == 'accounts_coordinators':
        # If the user is an accounts coordinator, retrieve the course of the user
        db_cursor.execute("SELECT course FROM accounts_coordinators WHERE username = %s", (username,))
        user_course = db_cursor.fetchone()

        if user_course:
            user_course = user_course[0]  # Extract the course from the result

            # Query reports where the course matches the user's course
            db_cursor.execute("SELECT * FROM forms_osd WHERE course = %s", (user_course,))
            reports = db_cursor.fetchall()
    else:
        # For other roles, simply retrieve reports for the logged-in user
        db_cursor.execute("SELECT * FROM forms_osd WHERE username = %s", (username,))
        reports = db_cursor.fetchall()
        user_course = ""

    # Close the cursor
    db_cursor.close()

    return render_template('forms.html', reports=reports, user_source=user_source, user_course=user_course)

@app.route('/download_form/<int:form_id>')
def download_form(form_id):
    db_cursor = db_connection.cursor()

    # Retrieve the file data for the given form_id from your database
    db_cursor.execute("SELECT filename, file_data FROM files WHERE id = %s", (form_id,))
    result = db_cursor.fetchone()

    if result is not None:
        filename, file_data = result

        # Serve the file as a downloadable attachment
        response = send_file(
            io.BytesIO(file_data),
            as_attachment=True,
            mimetype='application/pdf',
            download_name=filename + '.pdf'
            )
    
        return response

    # Handle the case where the form_id is not found
    return "Form not found", 404


@app.route('/hello', methods=['GET', 'POST'])
def homepage():
    username = session.get('username', '')
    
    if request.method == 'POST':
        # Handle the POST request (form submission)
        username = request.form['username']
        # Save the username in the session
        session['username'] = username

        

    # Determine the user source (accounts_cics or accounts_coordinators) and set the user_source variable
   

    db_cursor = db_connection.cursor()
    db_cursor.execute("SELECT * FROM accounts_cics WHERE username = %s", (username,))
    result_cics = db_cursor.fetchone()

    db_cursor.execute("SELECT * FROM accounts_coordinators WHERE username = %s", (username,))
    result_coordinators = db_cursor.fetchone()

    if result_cics:
        user_source = 'accounts_cics'
        session['source'] = user_source
    elif result_coordinators:
        user_source = 'accounts_coordinators'
        session['source'] = user_source
    else:
        user_source = 'unknown'  # Handle the case where the user source is not found

    # Close the cursor
    db_cursor.close()

    # Retrieve the profile picture path, name, and course for the logged-in user from the database
    db_cursor1 = db_connection.cursor()

    if user_source == 'accounts_cics':
        db_cursor1.execute("SELECT image_data, Name, CourseOrPosition FROM accounts_cics WHERE username = %s", (username,))
    elif user_source == 'accounts_coordinators':
        db_cursor1.execute("SELECT image_data, Name, Course FROM accounts_coordinators WHERE username = %s", (username,))
    else:
        # Handle the case where user_source is unknown
        db_cursor1.execute("SELECT image_data, Name, CourseOrPosition FROM accounts_cics WHERE username = %s", (username,))

    result_user_data = db_cursor1.fetchone()

    if result_user_data:
        profile_picture_data, name, course = result_user_data
        session['namestudent'] = name
        print(name)
    else:
        # Handle the case where user data is not found
        profile_picture_data = None
        name = "Name not found"
        course = "Course/Position not found"

    # Encode the profile picture data as a Base64 string
    if profile_picture_data is not None:
        profile_picture_base64 = base64.b64encode(profile_picture_data).decode('utf-8')
    else:
        profile_picture_base64 = None  # Handle the case where there is no profile picture data

    # Pass the sorted offenses, username, profile picture (Base64), name, course, and user_source to the template
    return render_template('homepage.html', username=username,profile_picture_base64=profile_picture_base64, name=name, course=course, user_source=user_source)

def lookup_student_info(username):
    try:
        db_cursor = db_connection.cursor(dictionary=True)
       

        # Assuming you have a table called 'students' with columns 'username', 'name', and 'course'
        query = "SELECT Name, CourseOrPosition FROM accounts_cics WHERE username = %s"
        db_cursor.execute(query, (username,))
        student_data = db_cursor.fetchone()

        if student_data:
            student_name = student_data['Name']
            student_course = student_data['CourseOrPosition']
            return student_name, student_course
        else:
            # Return None if the student is not found
            return None, None
    except mysql.connector.Error as err:
        # Handle any errors that may occur during the database query
        print(f"Error: {err}")
        return None, None
    finally:
        db_cursor.close()

# Usage example:
@app.route('/lookup_student', methods=['POST'])
def lookup_student():
    # Get the username from the request
    username = request.form.get('username')

    # Call the function to look up the student's name and course
    student_name, student_course = lookup_student_info(username)
    print(f"Student Name: {student_name}")
    print(f"Student Course: {student_course}")

    session['name'] =  student_name
    session['course'] = student_course

    # Return the result as JSON
    student_data = {'Name': student_name, 'CourseOrPosition': student_course}
    return jsonify(student_data)

@app.route('/download_report_file/<string:report_id>')
def download_report_file(report_id):
    db_cursor = db_connection.cursor()
    db_cursor.execute("SELECT file_form, file_form_name FROM reports WHERE report_id = %s", (report_id,))
    result = db_cursor.fetchone()

    if result is not None:
        file_data, file_name = result

        # Set the content type header to PDF
        content_type = 'application/pdf'

        # Set the filename to "default.pdf"
        response = make_response(file_data)
        response.headers['Content-Type'] = content_type

        # Set the filename to "default.pdf"
        response.headers['Content-Disposition'] = f'attachment; filename="{file_name}.pdf"'

        # Close the cursor after fetching the result
        db_cursor.close()

        return response

    # Handle the case where the file is not found
    db_cursor.close()
    return "File not found", 404



@app.route('/download_supporting_document/<string:report_id>')
def download_supporting_document(report_id):
    db_cursor = db_connection.cursor()
    db_cursor.execute("SELECT file_support_name, file_support_type, file_support FROM reports WHERE report_id = %s", (report_id,))
    result = db_cursor.fetchone()

    if result is not None:
        file_name, file_type, file_data = result

        # Set the content type header based on the supporting document's type stored in the database
        content_type = file_type

        # Set the filename to have the original name and extension
        response = make_response(file_data)
        response.headers['Content-Type'] = content_type

        # Set the filename based on the stored name and extension
        response.headers['Content-Disposition'] = f'attachment; filename="{file_name}{file_type}"'

        # Close the cursor after fetching the result
        db_cursor.close()

        return response

    # Handle the case where the file is not found
    db_cursor.close()
    return "File not found", 404

@app.route('/download_report_file1/<string:report_id>')
def download_report_file1(report_id):
    db_cursor = db_connection.cursor()
    db_cursor.execute("SELECT file_form, file_form_name FROM forms_osd WHERE form_id = %s", (report_id,))
    result = db_cursor.fetchone()

    if result is not None:
        file_data, file_name = result

        # Set the content type header to PDF
        content_type = 'application/pdf'

        # Set the filename to "default.pdf"
        response = make_response(file_data)
        response.headers['Content-Type'] = content_type

        # Set the filename to "default.pdf"
        response.headers['Content-Disposition'] = f'attachment; filename="{file_name}.pdf"'

        # Close the cursor after fetching the result
        db_cursor.close()

        return response

    # Handle the case where the file is not found
    db_cursor.close()
    return "File not found", 404



@app.route('/download_supporting_document1/<string:report_id>')
def download_supporting_document1(report_id):
    db_cursor = db_connection.cursor()
    db_cursor.execute("SELECT file_support_name, file_support_type, file_support FROM forms_osd WHERE form_id = %s", (report_id,))
    result = db_cursor.fetchone()

    if result is not None:
        file_name, file_type, file_data = result

        # Set the content type header based on the supporting document's type stored in the database
        content_type = file_type

        # Set the filename to have the original name and extension
        response = make_response(file_data)
        response.headers['Content-Type'] = content_type

        # Set the filename based on the stored name and extension
        response.headers['Content-Disposition'] = f'attachment; filename="{file_name}{file_type}"'

        # Close the cursor after fetching the result
        db_cursor.close()

        return response

    # Handle the case where the file is not found
    db_cursor.close()
    return "File not found", 404

@app.route('/change_report_status/<string:report_id>', methods=['POST'])
def change_report_status(report_id):
    new_status = request.form['new_status']
    print(new_status)
    print(report_id)
    db_cursor = db_connection.cursor()
    db_cursor.execute("UPDATE reports SET status = %s WHERE report_id = %s;", (new_status, report_id))
    db_connection.commit()  # Make sure to commit the changes to the database
    db_cursor.close()

    flash('Status has been successfully changed', 'success')

    return redirect(url_for('menu'))

@app.route('/change_report_status1/<string:report_id>', methods=['POST'])
def change_report_status1(report_id):
    new_status = request.form['new_status']
    print(new_status)
    print(report_id)
    db_cursor = db_connection.cursor()
    db_cursor.execute("UPDATE forms_osd SET status = %s WHERE form_id = %s;", (new_status, report_id))
    db_connection.commit()  # Make sure to commit the changes to the database
    db_cursor.close()

    flash('Status has been successfully changed', 'success')

    return redirect(url_for('requestpage'))

@app.route('/delete_report/<string:report_id>', methods=['POST'])
def delete_report(report_id):
    db_cursor = db_connection.cursor()
    db_cursor.execute("DELETE FROM reports WHERE report_id = %s;", (report_id,))
    db_connection.commit()  # Make sure to commit the changes to the database
    db_cursor.close()

    return redirect(url_for('menu'))

@app.route('/delete_report1/<string:report_id>', methods=['POST'])
def delete_report1(report_id):
    db_cursor = db_connection.cursor()
    db_cursor.execute("DELETE FROM forms_osd WHERE form_id = %s;", (report_id,))
    db_connection.commit()  # Make sure to commit the changes to the database
    db_cursor.close()

    return redirect(url_for('requestpage'))

@app.route('/delete_all_report1/<string:report_id>/<string:status>', methods=['POST'])
def delete_all_report1(report_id, status):
    if status == "Result":
        db_cursor = db_connection.cursor()
        db_cursor.execute("DELETE FROM forms_osd WHERE course = %s AND status = 'Approved' OR status = 'Rejected';", (report_id,))
        db_connection.commit()  # Make sure to commit the changes to the database
        db_cursor.close()
        return redirect(url_for('requestpage'))
        
    else:
        db_cursor = db_connection.cursor()
        db_cursor.execute("DELETE FROM forms_osd WHERE course = %s AND status = %s;", (report_id, status))
        db_connection.commit()  # Make sure to commit the changes to the database
        db_cursor.close()
        return redirect(url_for('requestpage'))

    

@app.route('/delete_all_report/<string:report_id>', methods=['POST'])
def delete_all_report(report_id):
    db_cursor = db_connection.cursor()
    db_cursor.execute("DELETE FROM reports WHERE course = %s;", (report_id,))
    db_connection.commit()  # Make sure to commit the changes to the database
    db_cursor.close()

    return redirect(url_for('menu'))

@app.route('/lookup_sanctions', methods=['POST'])
def lookup_sanctions():
    if request.method == 'POST':
        name = session.get('name', '')  # Updated to match the input name

        # Perform a database query to search for sanctions based on the username
        db_cursor = db_connection.cursor()
        db_cursor.execute("SELECT date_time, sanction FROM sanctions WHERE Username LIKE %s", ('%' + name + '%',))
        search_sanctions = db_cursor.fetchall()
        db_cursor.close()

        # Check if any sanctions were found
        if search_sanctions:
            # Convert datetime objects to string representations
            formatted_sanctions = [{'date_time': str(entry[0]), 'sanction': entry[1]} for entry in search_sanctions]
            return jsonify({'sanctions': formatted_sanctions})
        else:
            return jsonify({'error': 'No sanctions found'})
        
@app.route('/logout', methods=['GET'])
def logout():
    # Clear the session data
    session.clear()
    
    # Redirect the user to the login page or any other appropriate page
    return redirect('/')

@app.route('/fetch_sanctions', methods=['GET'])
def fetch_sanctions():
    student = session.get('namestudent', '')  # Retrieve the username from the session
    db_cursor = db_connection.cursor()
    db_cursor.execute("SELECT date_time, sanction FROM sanctions WHERE username = %s", (student,))
    sanctions_data = db_cursor.fetchall()
    db_cursor.close()

    # Debugging: Print the fetched data
    print("Fetched Data:", sanctions_data)

    # Convert the data to JSON and return it
    try:
        json_sanctions = [{"date": str(row[0]), "sanction": row[1]} for row in sanctions_data]
        return jsonify(json_sanctions)
    except Exception as e:
        print("Error converting data to JSON:", str(e))
        return jsonify({"error": "An error occurred while processing the data."}), 500

if __name__ == '__main__':
    app.run(debug=True, host='0.0.0.0')